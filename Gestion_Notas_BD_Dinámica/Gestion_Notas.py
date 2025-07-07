# bloque import
import csv
import names
import random
import math
from docx import Document


# bloque funciones
def crearBDDinamica(pcantidad, pporcentaje, pcantidad2, pporcentaje2, anno, n1, n2, n3):
    """
    Función: función de procesamiento, recibe toda la información y la mete a la base de datos
    Entradas:
    - pcantidad(int): es la cantidad de estudiantes que se sacan del primer recurso
    - pporcentaje(int): es el porcentaje de los estudiantes del primer recurso que se agregan a la base
    - pcantidad2(int): es la cantidad de estudiantes que se sacan del segundo recurso
    - pporcentaje2(int): es el porcentaje de los estudiantes del segundo recurso que se agregan a la base
    - anno(int): el año en el que ingresó el estudiante
    - n1(int): la primera nota de evaluación metida por el usuario
    - n2(int): la segunda nota de evaluación metida por el usuario
    - n3(int): la tercera nota de evaluación metida por el usuario
    Salidas:
    """

    for i in range(
        round(obtenerCantidad(pcantidad, pporcentaje))
    ):  # se saca el total de estudiantes del primer recurso por su cantidad y porcentaje
        genero = "Masculino"
        if (
            random.randint(0, 1) == 1
        ):  # este es el género, 0 es masculino y 1 es femenino
            genero = "Femenino"

        nombre = names.get_first_name()
        apellido = names.get_last_name()
        carne = generarCarne(anno)
        nuevoNombre = [
            nombre,
            apellido,
            names.get_last_name(),
            genero,
            carne,
            generarCorreo(nombre, apellido, carne),
            generarNota(n1, n2, n3),
        ]
        with open(
            "BDDinamica.csv", "a", newline=""
        ) as archivo:  # esta estrategia fue sacada de https://www.youtube.com/watch?v=KKsy1Mdrgq4
            writer = csv.writer(archivo)
            writer.writerow(nuevoNombre)

    # este with saca los nombres de el archivo estudiantes.txt
    with open(
        "estudiantes.txt", "r", encoding="utf-8"
    ) as archivoEstudiantes:  # fue añádido el encoding="utf-8" porque no se mostraban bien los carácterez especiales
        lineas = archivoEstudiantes.readlines()

    # este with mete los estudiantes del primer archivo a la base de datos
    with open("BDDinamica.csv", "a", newline="", encoding="utf-8") as archivo:
        writer = csv.writer(archivo)
        for i in range(
            min(round(obtenerCantidad(pcantidad2, pporcentaje2)), len(lineas))
        ):

            linea = lineas[i].strip().split(",")  # agrega el nombre, apellidos y género
            linea.append(generarCarne(anno))  # agrega el carné
            nombre = linea[0]
            apellido = linea[1]
            carne = linea[3]
            linea.append(generarCorreo(nombre, apellido, carne))  # agrega el correo
            linea.append(generarNota(n1, n2, n3))  # agrega la nota
            writer.writerow(linea)  # escribe lo agregrado en la base de datos
    print("Base De Datos Dinámica creada")


def obtenerCantidad(cant, porc):
    """
    Función: obtiene la cantidad de estudiantes en total
    Entradas:
    - cant(int) = es el número total de estudiantes que nos da el usuario
    - porc(int) = es el porcentaje de estudiantes que son elegidos
    Salidas:
    - total = es el número resultante de sacarle el porcentaje a el número
    """
    total = 0
    total = cant * (porc * 0.01)
    return round(total)


def generarCarne(anno):
    """
    Función: genera un carné aleatorio del estudiante
    Entradas:
    - anno(int) = es el año en el que ingresó el estudiante
    Salidas:
    - carne(str) = es el carné final del estudiante
    """
    carne = ""
    carne += (
        str(anno) + "0" + str(random.randint(1, 6)) + str(random.randint(1000, 9999))
    )
    return carne


def generarCorreo(nom, ape, car):
    """
    Función: genera un correo al estudiante
    Entradas:
    - nom(str) = es el nombre del estudiante
    - ape(str) = es el apellido del estudiante
    - car(str) = es el carné del estudiante
    Salidas:
    - correo(str) = es el correo del estudiante
    """
    correo = ""
    correo += (
        str(nom[:1]) + str(ape) + str(car[6:]) + "@estudiantec.cr"
    )  # arma el correo
    return (
        correo.lower()
    )  # lo retorna, pero los nombres salen en mayúscula, por eso es necesario poner el .lower()


def generarNota(n1, n2, n3):
    """
    Función: genera aleatoriamente las notas de los estudiantes, favoreciendo notas más altas.
    Entradas:
    - n1(int): valor porcentual de la primera evaluación
    - n2(int): valor porcentual de la segunda evaluación
    - n3(int): valor porcentual de la tercera evaluación
    Salidas:
    - tupla de 5 elementos: (nota1, nota2, nota3, total, totalR)
    """

    nota1 = round(n1 * (random.uniform(0.7, 1.0)), 1)  # 70% - 100% del valor de n1
    nota2 = round(n2 * (random.uniform(0.7, 1.0)), 1)
    nota3 = round(n3 * (random.uniform(0.7, 1.0)), 1)

    total = round(nota1 + nota2 + nota3, 1)
    return (nota1, nota2, nota3, total, total)  # nota final duplicada como en original



# entradar y salidas, y auxiliares
def crearBDDinamicaAux(
    pcantidad, pporcentaje, pcantidad2, pporcentaje2, anno, n1, n2, n3
):
    """
    Función: valora la validez de las variables ingresadas
    Entradas:
    - pcantidad(int): es la cantidad de estudiantes que se sacan del primer recurso
    - pporcentaje(int): es el porcentaje de los estudiantes del primer recurso que se agregan a la base
    - pcantidad2(int): es la cantidad de estudiantes que se sacan del segundo recurso
    - pporcentaje2(int): es el porcentaje de los estudiantes del segundo recurso que se agregan a la base
    - anno(int): el año en el que ingresó el estudiante
    - n1(int): la primera nota de evaluación metida por el usuario
    - n2(int): la segunda nota de evaluación metida por el usuario
    - n3(int): la tercera nota de evaluación metida por el usuario
    Salidas:
    - las mismas variables ya valoradas
    """

    try:
        pcantidad = int(pcantidad)
        pporcentaje = int(pporcentaje)
        pcantidad2 = int(pcantidad2)
        pporcentaje2 = int(pporcentaje2)
        n1 = int(n1)
        n2 = int(n2)
        n3 = int(n3)
        if pcantidad > 0 and pporcentaje > 0:
            if n1 + n2 + n3 == 100:
                return crearBDDinamica(
                    pcantidad, pporcentaje, pcantidad2, pporcentaje2, anno, n1, n2, n3
                )
            else:
                print("Los valores de las evaluaciones deben sumar 100 entre los tres.")
                return crearBDDinamicaES()
        else:
            return crearBDDinamicaES()
            print("Ingrese datos validos.")

    except ValueError:
        print("Ingrese datos válidos.")
        return crearBDDinamicaES()


def crearBDDinamicaES():
    """
    Función: Pide los datos al usuario
    Entradas: N/A
    Salidas:
    - pcantidad(int): es la cantidad de estudiantes que se sacan del primer recurso
    - pporcentaje(int): es el porcentaje de los estudiantes del primer recurso que se agregan a la base
    - pcantidad2(int): es la cantidad de estudiantes que se sacan del segundo recurso
    - pporcentaje2(int): es el porcentaje de los estudiantes del segundo recurso que se agregan a la base
    - anno(int): el año en el que ingresó el estudiante
    - n1(int): la primera nota de evaluación metida por el usuario
    - n2(int): la segunda nota de evaluación metida por el usuario
    - n3(int): la tercera nota de evaluación metida por el usuario
    """
    # cantidad1 = input(
    #     "Ingrese la cantidad de alumnos que se sacarán del primer recurso: "
    # )
    # porcentaje1 = input(
    #     "Ingrese el porcentaje de alumnos que se sacarán del primer recurso: "
    # )
    # cantidad2 = input(
    #     "Ingrese la cantidad de alumnos que se sacarán del segundo recurso: "
    # )
    # porcentaje2 = input(
    #     "Ingrese el porcentaje de alumnos que se sacarán del segundo recurso: "
    # )
    # anno = input("Ingrese el año en el que ingresó: ")

    cantidad1 = 120
    porcentaje1 = 45
    cantidad2 = 75
    porcentaje2 = 45
    anno = 2020

    # n1 = input("Ingrese el valor de la primera evaluación: ")
    # n2 = input("Ingrese el valor de la segunda evaluación: ")
    # n3 = input("Ingrese el valor de la tercera evaluación: ")
    n1 = 25
    n2 = 35
    n3 = 40

    return crearBDDinamicaAux(
        cantidad1, porcentaje1, cantidad2, porcentaje2, anno, n1, n2, n3
    )

# --------------------------------------------------------------------------------------------------------------------------------------------------
def registrarEstudiante():
    """ 
    Función que registra un nuevo estudiante solicitando datos personales, notas y generando carné, correo y nota final.
    Entradas:
        - Datos por input: nombre, apellidos, género, año de ingreso, porcentaje de cada evaluación (N1, N2, N3).
    Salidas: 
        - Archivo 'BDDinamica.csv' actualizado con los datos del estudiante.
        - Mensaje de confirmación por consola.
    """

    nombre = input("Nombre: ")
    apellido1 = input("Primer apellido: ")
    apellido2 = input("Segundo apellido: ")
    genero = input("Género (Masculino/Femenino): ")
    anno = int(input("Año de ingreso: "))
    carne = generarCarne(anno)
    correo = generarCorreo(nombre, apellido1, carne)
    
    print("Ingrese el valor de las evaluaciones (suman 100 en total)")
    n1 = int(input("Nota 1 (porcentaje): "))
    n2 = int(input("Nota 2 (porcentaje): "))
    n3 = int(input("Nota 3 (porcentaje): "))

    if n1 + n2 + n3 != 100:
        print("Las evaluaciones deben sumar 100.")
        return

    nota = generarNota(n1, n2, n3)

    estudiante = [nombre, apellido1, apellido2, genero, carne, correo, nota]

    with open("BDDinamica.csv", "a", newline="", encoding="utf-8") as archivo:
        writer = csv.writer(archivo)
        writer.writerow(estudiante)

    print("Estudiante registrado correctamente.")

def reporteGenero():
    """ 
    Función que genera reportes en archivos Word separados por género con el cálculo de la nota de acta de cada estudiante.
    Entradas:
        - Archivo 'BDDinamica.csv'.
    Salidas: 
        - Archivos 'mujeres.docx' y 'hombres.docx' con los reportes.
        - Mensaje de confirmación por consola para cada archivo generado.
    """

    # Porcentajes definidos para cada nota
    porcentaje_n1 = 0.30
    porcentaje_n2 = 0.30
    porcentaje_n3 = 0.40

    estudiantes = {"Femenino": [], "Masculino": []}

    # Leer el archivo CSV
    with open("BDDinamica.csv", "r", encoding="utf-8") as archivo:
        reader = csv.reader(archivo)
        for fila in reader:
            try:
                # Reconstruir nombre completo
                nombre = f"{fila[0]} {fila[1]} {fila[2]}"
                genero = fila[3]
                carne = fila[4]
                correo = fila[5]

                # Extraer las notas desde la cadena "(n1, n2, n3, curva, final)"
                notas_str = fila[6].strip("()").split(",")
                if len(notas_str) < 3:
                    raise ValueError("Faltan notas")

                # Convertir las tres primeras a float
                n1 = float(notas_str[0].strip())
                n2 = float(notas_str[1].strip())
                n3 = float(notas_str[2].strip())

                nota_acta = round(n1 * porcentaje_n1 + n2 * porcentaje_n2 + n3 * porcentaje_n3, 1)

                estudiante = {
                    "nota_acta": nota_acta,
                    "n1": n1,
                    "n2": n2,
                    "n3": n3,
                    "nombre": nombre,
                    "carne": carne,
                    "correo": correo
                }

                if genero in estudiantes:
                    estudiantes[genero].append(estudiante)
            except Exception as e:
                print(f"Error en fila: {fila} -> {e}")
                continue

    # Función para crear el Word por género
    def crear_doc(estudiantes_genero, genero):
        """ 
        Función que genera un documento Word (.docx) con el reporte de estudiantes de un género específico, 
        ordenados por su nota de acta de mayor a menor. Incluye información personal y notas ponderadas.

        Entradas:
            - estudiantes_genero (list): Lista de diccionarios con los datos de cada estudiante.
            - genero (str): Género de los estudiantes ("Femenino" o "Masculino").

        Salidas:
            - Archivo Word guardado como 'mujeres.docx' o 'hombres.docx' con el listado de estudiantes.
            - Mensaje por consola confirmando la generación del archivo.

        Método tomado y adaptado de: 
        https://python--docx-readthedocs-io.translate.goog/en/latest/?_x_tr_sl=en&_x_tr_tl=es&_x_tr_hl=es&_x_tr_pto=tc
        """
        doc = Document()
        doc.add_heading(f"Reporte de estudiantes - {genero}", 0)

        # Ordenar por nota de acta de mayor a menor
        estudiantes_genero.sort(key=lambda x: x["nota_acta"], reverse=True)

        for est in estudiantes_genero:
            linea = f"{est['nota_acta']} | {est['n1']}, {est['n2']}, {est['n3']} | {est['nombre']} | {est['carne']} | {est['correo']}"
            doc.add_paragraph(linea)

        doc.add_paragraph("")
        doc.add_paragraph(f"Porcentajes aplicados: N1 = {int(porcentaje_n1*100)}%, N2 = {int(porcentaje_n2*100)}%, N3 = {int(porcentaje_n3*100)}%")
        doc.add_paragraph(f"Total de estudiantes ({genero}): {len(estudiantes_genero)}")

        nombre_archivo = "mujeres.docx" if genero == "Femenino" else "hombres.docx"
        doc.save(nombre_archivo)
        print(f"Archivo '{nombre_archivo}' generado.")

    # Crear los documentos
    crear_doc(estudiantes["Femenino"], "Femenino")
    crear_doc(estudiantes["Masculino"], "Masculino")

def enviarCorreosReposicion():
    """ 
    Función que simula el envío de correos a estudiantes con nota de acta menor a 70, indicando que deben presentarse a reposición.
    Entradas:
        - Archivo 'BDDinamica.csv'.
    Salidas: 
        - Lista por consola con los nombres, correos y notas de los estudiantes que deben hacer reposición.
    """

    print("\nCorreos enviados a estudiantes con nota de acta menor a 70:")

    with open("BDDinamica.csv", "r", encoding="utf-8") as archivo:
        reader = csv.reader(archivo)
        for fila in reader:
            try:
                nombre = f"{fila[0]} {fila[1]} {fila[2]}"
                correo = fila[5]
                notas_str = fila[6].strip("()")
                notas = [float(n.strip()) for n in notas_str.split(",")]
                nota_acta = notas[-1]

                if nota_acta < 70:
                    print(f"- {nombre}: {correo} (Nota: {nota_acta})")
            except Exception as e:
                print(f"Error procesando fila {fila}: {e}")

def estadisticaGeneracion():
    """ 
    Función que calcula y muestra estadísticas de aprobación, reposición y reprobación por año de ingreso (generación).
    Entradas:
        - Archivo 'BDDinamica.csv'.
    Salidas: 
        - Estadísticas por consola agrupadas por año de ingreso.
    """

    generaciones = {}

    with open("BDDinamica.csv", "r", encoding="utf-8") as archivo:
        reader = csv.reader(archivo)
        for fila in reader:
            try:
                carne = fila[4]
                anno = carne[:4]

                notas_str = fila[6].strip("()")
                notas = [float(n.strip()) for n in notas_str.split(",")]
                nota_acta = notas[-1]

                # Inicializar si no existe
                if anno not in generaciones:
                    generaciones[anno] = {"Aprobados": 0, "Reposicion": 0, "Reprobados": 0, "Totales": 0}

                # Clasificación
                if nota_acta >= 70:
                    generaciones[anno]["Aprobados"] += 1
                elif nota_acta >= 50:
                    generaciones[anno]["Reposicion"] += 1
                else:
                    generaciones[anno]["Reprobados"] += 1

                generaciones[anno]["Totales"] += 1

            except Exception as e:
                print(f"Error en fila: {fila} -> {e}")

    # Mostrar resultados
    print("\nEstadística por generación:")
    print(f"{'Año':<6}{'Aprobados':>10}{'Reposición':>12}{'Reprobados':>12}{'Totales':>10}")
    
    total_ap = total_rep = total_reprob = total_gen = 0

    for anno in sorted(generaciones.keys()):
        datos = generaciones[anno]
        print(f"{anno:<6}{datos['Aprobados']:>10}{datos['Reposicion']:>12}{datos['Reprobados']:>12}{datos['Totales']:>10}")
        total_ap += datos['Aprobados']
        total_rep += datos['Reposicion']
        total_reprob += datos['Reprobados']
        total_gen += datos['Totales']

    print(f"{'Totales':<6}{total_ap:>10}{total_rep:>12}{total_reprob:>12}{total_gen:>10}")


def reporteBuenRendimientoPorSede():
    """ 
    Función que permite consultar estudiantes con buen rendimiento (nota final >= 70) por sede, identificada por el carné.
    Entradas:
        - Archivo 'BDDinamica.csv'.
        - Input de usuario para seleccionar la sede a consultar.
    Salidas: 
        - Lista por consola de estudiantes con buen rendimiento en la sede seleccionada.
    """

    sedes = {str(i): [] for i in range(1, 7)}

    with open("BDDinamica.csv", "r", encoding="utf-8") as archivo:
        reader = csv.reader(archivo)
        for fila in reader:
            if len(fila) < 7:
                continue  # fila incompleta

            carne = fila[4].strip()

            if len(carne) < 6:
                continue  # carné inválido

            sede = carne[5]  # sexto carácter

            if sede not in sedes:
                continue  # sede no válida (no está entre "1"-"6")

            try:
                notas = eval(fila[6])
                if len(notas) >= 4 and notas[3] >= 70:  # usa nota final
                    sedes[sede].append(fila)
            except:
                continue  # error en formato de notas

    print("\nSedes disponibles:")
    for sede in sorted(sedes.keys()):
        print(f"{sede}. Sede {sede} ({len(sedes[sede])} con buen rendimiento)")

    seleccion = input("Ingrese el número de sede que desea consultar (1-6): ").strip()

    if seleccion in sedes and sedes[seleccion]:
        print(f"\nEstudiantes con buen rendimiento en Sede {seleccion}:")
        for i, estudiante in enumerate(sedes[seleccion], start=1):
            notas = eval(estudiante[6])
            print(f"{i}. {estudiante[0]} {estudiante[1]} - {estudiante[4]} - Nota final: {notas[3]}")
    else:
        print("No hay estudiantes con buen rendimiento en esta sede.")


# programa principal

def menuPrincipal():
    while True:
        print("\n--- MENÚ PRINCIPAL ---")
        print("1.  Crear Base de Datos Dinámica")
        print("2.  Registrar Estudiante Manualmente")
        print("3.  Generar reporte HTML y .csv")
        print("4.  Respaldar en XML")
        print("5.  Generar Reporte por Género (Word)")
        print("6.  Gestionar curva")
        print("7.  Enviar Correos de Reposición (mostrar en consola)")
        print("8.  Aplazados en al menos 2 exámenes (.pdf)")
        print("9.  Estadística por generación")
        print("10. Reporte por sede con buen rendimiento.")
        print("11. Salir")

        opcion = input("Seleccione una opción: ")

        if opcion == "1":
            crearBDDinamicaES()
        elif opcion == "2":
            registrarEstudiante()
        elif opcion == "3":
            print("Falta")
        elif opcion == "4":
            print("Falta")
        elif opcion == "5":
            reporteGenero()
        elif opcion == "6":
            print("Falta")
        elif opcion == "7":
            enviarCorreosReposicion()
        elif opcion == "8":
            print("Falta")
        elif opcion == "9":
            estadisticaGeneracion()
        elif opcion == "10":
            reporteBuenRendimientoPorSede()
        elif opcion == "11":
            print("¡Hasta luego!")
            break
        else:
            print("Opción no válida. Intente nuevamente.")

# Ejecutar el menú
menuPrincipal()
